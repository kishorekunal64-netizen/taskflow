"""
generate_sop_docx.py — Convert RAGAI_SOP.md to a formatted Word document.

Usage:
    venv\\Scripts\\activate
    python generate_sop_docx.py

Output: RAGAI_SOP.docx in the project root.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc: Document) -> None:
    """Add a thin horizontal line (paragraph border)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def style_inline(para, text: str) -> None:
    """Add a paragraph run, handling **bold** and `code` inline markers."""
    # Split on **bold** and `code` markers
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            if part:
                para.add_run(part)


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

def convert(md_path: Path, docx_path: Path) -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    i = 0
    in_code_block = False
    code_lines: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return

        # Filter out separator rows (---|---)
        data = [r for r in table_rows if not all(re.match(r"^[-: ]+$", c) for c in r)]
        if not data:
            in_table = False
            table_rows = []
            return

        col_count = max(len(r) for r in data)
        # Pad rows
        data = [r + [""] * (col_count - len(r)) for r in data]

        tbl = doc.add_table(rows=len(data), cols=col_count)
        tbl.style = "Table Grid"

        for ri, row in enumerate(data):
            for ci, cell_text in enumerate(row):
                cell = tbl.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                style_inline(p, cell_text.strip())
                for run in p.runs:
                    run.font.size = Pt(9.5)
                if ri == 0:
                    set_cell_bg(cell, "1F3864")
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
                elif ri % 2 == 0:
                    set_cell_bg(cell, "EEF2F7")

        doc.add_paragraph()
        in_table = False
        table_rows = []

    while i < len(lines):
        line = lines[i]

        # ---- Code block ----
        if line.strip().startswith("```"):
            if not in_code_block:
                if in_table:
                    flush_table()
                in_code_block = True
                code_lines = []
            else:
                # End of code block — render
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    run = p.add_run("\n".join(code_lines))
                    run.font.name = "Courier New"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
                    # Light grey shading on the paragraph
                    pPr = p._p.get_or_add_pPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "F3F3F3")
                    pPr.append(shd)
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ---- Table row ----
        if line.strip().startswith("|"):
            if in_table is False:
                in_table = True
                table_rows = []
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # ---- Horizontal rule ----
        if re.match(r"^---+$", line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ---- Blank line ----
        if not line.strip():
            i += 1
            continue

        # ---- Headings ----
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            # Strip markdown link anchors like (#section)
            text = re.sub(r"\(#[^)]+\)", "", text).strip()
            heading_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
            h = doc.add_heading(text, level=level)
            h.style = doc.styles[heading_map.get(level, "Heading 4")]
            # Colour heading 1 dark blue, others slightly lighter
            for run in h.runs:
                if level == 1:
                    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                elif level == 2:
                    run.font.color.rgb = RGBColor(0x2E, 0x54, 0x96)
                else:
                    run.font.color.rgb = RGBColor(0x1F, 0x3D, 0x7A)
            i += 1
            continue

        # ---- Blockquote ----
        if line.startswith(">"):
            text = line.lstrip("> ").strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ---- Bullet list ----
        if re.match(r"^(\s*[-*])\s+", line):
            indent = len(line) - len(line.lstrip())
            text = re.sub(r"^\s*[-*]\s+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.25 + indent * 0.02)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            style_inline(p, text)
            i += 1
            continue

        # ---- Numbered list ----
        if re.match(r"^\d+\.\s+", line):
            text = re.sub(r"^\d+\.\s+", "", line)
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            style_inline(p, text)
            i += 1
            continue

        # ---- Checkbox list ----
        if re.match(r"^- \[[ x]\]", line):
            checked = "[x]" in line
            text = re.sub(r"^- \[[ x]\]\s*", "", line)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            mark = "☑ " if checked else "☐ "
            run = p.add_run(mark)
            run.font.color.rgb = RGBColor(0x2E, 0x54, 0x96)
            style_inline(p, text)
            i += 1
            continue

        # ---- Normal paragraph ----
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        style_inline(p, line.strip())
        i += 1

    if in_table:
        flush_table()

    doc.save(str(docx_path))
    print(f"✅  Saved: {docx_path.resolve()}")


if __name__ == "__main__":
    convert(Path("RAGAI_SOP.md"), Path("RAGAI_SOP.docx"))
